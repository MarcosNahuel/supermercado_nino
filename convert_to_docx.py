from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_break(doc):
    """Agrega un salto de página"""
    doc.add_page_break()

def add_horizontal_line(paragraph):
    """Agrega una línea horizontal"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E5090')
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_table_border(table):
    """Establece bordes de tabla"""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '2E5090')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def create_styled_document():
    """Crea el documento con estilos estéticos"""

    # Leer el archivo markdown
    with open(r'd:\OneDrive\GitHub\supermercado_nino\outputs\ENTREGABLE_FASE1_DON_NINO_V2.md', 'r', encoding='utf-8') as f:
        content = f.read()

    # Crear documento
    doc = Document()

    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = content.split('\n')
    i = 0
    in_table = False
    table_data = []

    while i < len(lines):
        line = lines[i].strip()

        # Título principal (# )
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            p = doc.add_heading(text, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(28)
                run.font.color.rgb = RGBColor(46, 80, 144)
                run.font.bold = True

        # Subtítulos nivel 2 (## )
        elif line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_heading(text, level=1)
            for run in p.runs:
                run.font.size = Pt(20)
                run.font.color.rgb = RGBColor(46, 80, 144)
                run.font.bold = True
            # Agregar línea decorativa
            para = doc.add_paragraph()
            add_horizontal_line(para)

        # Subtítulos nivel 3 (### )
        elif line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_heading(text, level=2)
            for run in p.runs:
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(70, 120, 180)
                run.font.bold = True

        # Subtítulos nivel 4 (#### )
        elif line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_heading(text, level=3)
            for run in p.runs:
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(100, 140, 200)
                run.font.bold = True

        # Líneas horizontales (---)
        elif line.startswith('---'):
            para = doc.add_paragraph()
            add_horizontal_line(para)

        # Citas (> )
        elif line.startswith('> '):
            text = line[2:].strip()
            p = doc.add_paragraph(text)
            p.style = 'Intense Quote'
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.italic = True
                run.font.color.rgb = RGBColor(70, 70, 70)

        # Detectar inicio de tabla
        elif '|' in line and not in_table:
            in_table = True
            table_data = []
            # Leer todas las líneas de la tabla
            while i < len(lines) and '|' in lines[i]:
                table_line = lines[i].strip()
                if table_line and not table_line.replace('|', '').replace('-', '').strip() == '':
                    if not all(c in '|-: ' for c in table_line):
                        cells = [cell.strip() for cell in table_line.split('|')]
                        cells = [c for c in cells if c]
                        table_data.append(cells)
                i += 1

            # Crear tabla
            if table_data:
                num_cols = len(table_data[0])
                num_rows = len(table_data)
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Light Grid Accent 1'

                # Llenar tabla
                for row_idx, row_data in enumerate(table_data):
                    for col_idx, cell_data in enumerate(row_data):
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = cell_data

                        # Formato de encabezado
                        if row_idx == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(255, 255, 255)
                            cell._element.get_or_add_tcPr().append(
                                OxmlElement('w:shd')
                            )
                            cell._element.get_or_add_tcPr().find(
                                qn('w:shd')
                            ).set(qn('w:fill'), '2E5090')

                        # Formato general
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(10)

                set_table_border(table)
                doc.add_paragraph()  # Espacio después de la tabla

            in_table = False
            continue

        # Texto en negrita (**texto**)
        elif '**' in line:
            p = doc.add_paragraph()
            parts = line.split('**')
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                if idx % 2 == 1:  # Texto entre **
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(46, 80, 144)
                run.font.size = Pt(11)

        # Lista con viñetas (- )
        elif line.startswith('- '):
            text = line[2:].strip()
            # Procesar texto con formato
            p = doc.add_paragraph(style='List Bullet')
            if '**' in text:
                parts = text.split('**')
                for idx, part in enumerate(parts):
                    run = p.add_run(part)
                    if idx % 2 == 1:
                        run.font.bold = True
                    run.font.size = Pt(11)
            else:
                run = p.add_run(text)
                run.font.size = Pt(11)

        # Texto normal
        elif line and not in_table:
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(11)

        i += 1

    # Guardar documento
    doc.save(r'd:\OneDrive\GitHub\supermercado_nino\outputs\ENTREGABLE_FASE1_DON_NINO_V2.docx')
    print("Documento Word creado exitosamente: ENTREGABLE_FASE1_DON_NINO_V2.docx")

if __name__ == '__main__':
    create_styled_document()
